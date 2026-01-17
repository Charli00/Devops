from docx import Document

# Create a new Word document for the 1-page ATS resume
doc = Document()

# Header
doc.add_heading("PENIEL CHARLES P", level=1)
doc.add_paragraph("Tirunelveli, Tamil Nadu, Phone: +91 9741398794, Email: penialcharles007@gmail.com")
doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/peniel-charles-394b46229")

# Professional Summary
doc.add_heading("PROFESSIONAL SUMMARY", level=2)
summary_points = [
    "Oracle Database Administrator with 3 years of experience in managing production, development, and UAT environments.",
    "Proficient in Oracle RAC, ASM, RMAN, AWR analysis, Data Pump, and performance tuning across Oracle 12c/19c.",
    "Skilled in automation, storage management, and OCI-based disaster recovery for high-availability enterprise systems.",
    "Committed to optimizing database reliability, uptime (>99.9%), and operational efficiency in mission-critical setups."
]
for point in summary_points:
    doc.add_paragraph(point, style="List Bullet")

# Technical Skills & Core Competencies
doc.add_heading("TECHNICAL SKILLS & CORE COMPETENCIES", level=2)
skills = {
    "Databases": "Oracle 12c, 19c",
    "Backup & Recovery": "RMAN, Data Pump (expdp/impdp)",
    "Performance Tuning": "AWR/ASH reports, SQL tuning, query optimization",
    "Storage & ASM": "Disk group management, LUN integration, rebalance monitoring",
    "High Availability": "Oracle RAC, Data Guard (monitoring), OCI DR drills",
    "Automation": "Shell scripting, crontab scheduling, log management",
    "Patching & Maintenance": "Grid/DB home patching, pre & post validation",
    "Tools": "SQL*Plus, OEM, OCI Console, Crontab",
    "Operating Systems": "Linux (RHEL, CentOS), Windows Server"
}
for category, details in skills.items():
    doc.add_paragraph(f"{category}: {details}")

# Professional Experience
doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
p = doc.add_paragraph()
p.add_run("Oracle Database Administrator – 3i Infotech, Bengaluru, Karnataka\n").bold = True
p.add_run("March 2023 – Present").italic = True

experience_points = [
    "Delivered 24x7 on-call support for production and development databases, ensuring >99.9% uptime and smooth business operations.",
    "Automated maintenance tasks (index rebuilds, alert log checks) using Shell scripting, reducing manual effort by 40%.",
    "Configured and managed database security, roles, and privileges to eliminate access errors and ensure compliance.",
    "Performed RMAN and Data Pump backups/restores with zero data loss during migrations and environment refreshes.",
    "Analyzed AWR reports and tuned SQL queries to improve application response times by up to 25%.",
    "Supported Grid/DB home patching and post-validation using OCI console and Oracle Enterprise Manager.",
    "Managed GoldenGate replication, resolving lag issues and ensuring continuous data synchronization.",
    "Executed disaster recovery drills in OCI, validating high-availability procedures.",
    "Coordinated with storage and OS teams to allocate additional LUN space and expand ASM disk groups for capacity optimization.",
    "Collaborated on ASM storage cleanup, enhancing disk efficiency and performance stability."
]
for point in experience_points:
    doc.add_paragraph(point, style="List Bullet")

# Education
doc.add_heading("EDUCATION", level=2)
doc.add_paragraph("Bachelor of Technology (B.Tech) in Information Technology")
doc.add_paragraph("PSN Engineering College, Tirunelveli, Tamil Nadu — 7.5 CGPA")
doc.add_paragraph("March 2018 – August 2022")

# Save the document
output_path = "/workspaces/Devops/Peniel_Charles_Oracle_DBA_1Page_ATS_Resume.docx"
doc.save(output_path)

output_path
