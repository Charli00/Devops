from docx import Document

# Create the improved final resume document
doc = Document()

# Header
doc.add_heading("Peniel Charles P", level=1)

contact_info = doc.add_paragraph()
contact_info.add_run("Phone: +91 97413 98794 | Email: penialcharles007@gmail.com | ")
contact_info.add_run("LinkedIn: linkedin.com/in/peniel-charles-394b46229\n")
contact_info.add_run("Open to relocation (Anywhere in India)").italic = True

# Career Objective
doc.add_heading("Career Objective", level=2)
doc.add_paragraph(
    "Results-driven Oracle Database Administrator with 3 years of experience in Oracle 12c/19c RAC and ASM environments. "
    "Proficient in RMAN, Data Pump, AWR/ASH/ADDM tuning, and automation using Shell scripting. Skilled in performance tuning, "
    "disaster recovery, and Oracle Cloud Infrastructure (OCI) management. Seeking to leverage expertise in backup, recovery, "
    "and high availability to ensure database reliability and performance in enterprise environments."
)

# Key Highlights
doc.add_heading("Key Highlights", level=2)
doc.add_paragraph(
    "• 3 years of hands-on experience managing Oracle 12c/19c RAC and ASM databases.\n"
    "• Expertise in RMAN backup and recovery, AWR/ASH/ADDM performance tuning, and SQL optimization.\n"
    "• Automated DBA tasks using Shell scripting, reducing manual workload by 40%.\n"
    "• Strong knowledge of Oracle GoldenGate replication, patching, and OCI-based DR solutions.\n"
)

# Professional Experience
doc.add_heading("Professional Experience", level=2)
doc.add_paragraph().add_run("Oracle Database Administrator — 3i Infotech").bold = True
doc.add_paragraph("Bengaluru, Karnataka | Mar 2023 – Present").italic = True

experience = doc.add_paragraph()
experience.add_run(
    "• Administered Oracle 12c/19c databases in RAC and ASM environments ensuring 99.9% uptime across production and UAT systems.\n"
    "• Implemented RMAN full/incremental backups and Data Pump imports/exports; validated recovery procedures ensuring 100% data reliability.\n"
    "• Analyzed and tuned SQL queries using AWR, ASH, and ADDM reports, reducing query response time by 35%.\n"
    "• Automated maintenance tasks (index rebuilds, alert log cleanup, tablespace monitoring) using Shell scripts and Crontab scheduling.\n"
    "• Applied database and Grid Infrastructure patches (19c) with zero downtime through rolling patching on RAC nodes.\n"
    "• Managed Oracle GoldenGate replication across multiple environments; resolved lag and synchronization issues improving data latency by 60%.\n"
    "• Conducted Disaster Recovery (DR) drills using RMAN and OCI-based backups, ensuring business continuity and compliance.\n"
    "• Collaborated with infrastructure teams for storage optimization, capacity planning, and proactive issue management.\n"
    "• Supported user management, tablespace growth monitoring, and incident troubleshooting using OEM and SQL*Plus.\n"
)

# Key Projects
doc.add_heading("Key Projects", level=2)
projects = doc.add_paragraph()
projects.add_run("1. Automated Materialized View Refresh (Oracle 19c, Shell Scripting, Crontab)\n").bold = True
projects.add_run(
    "Developed a Shell-based automation framework to refresh and validate materialized views daily. Reduced refresh cycle time by 50% "
    "and eliminated manual intervention for 20+ MVs.\n"
)
projects.add_run("2. Database Migration via Data Pump & RMAN (Oracle 19c, OCI)\n").bold = True
projects.add_run(
    "Led migration of production schemas using Data Pump (expdp/impdp) and RMAN duplication techniques to OCI cloud servers. "
    "Achieved zero data loss and verified consistency through checksum validation.\n"
)
projects.add_run("3. Oracle Grid Patching & Performance Validation (19c RAC, AWR/ADDM)\n").bold = True
projects.add_run(
    "Applied Oracle 19c Grid patches across multi-node RAC clusters using out-of-place patching strategy. "
    "Validated patch success via AWR and ADDM analysis ensuring stable post-patch performance.\n"
)

# Technical Skills
doc.add_heading("Technical Skills", level=2)
skills = doc.add_paragraph()
skills.add_run("Database Administration: ").bold = True
skills.add_run("Oracle 12c, 19c, RAC, ASM, OEM, User Management, Data Security & Auditing\n")
skills.add_run("Backup & Recovery: ").bold = True
skills.add_run("RMAN, Data Pump (expdp/impdp), Incremental Backup, Recovery Validation\n")
skills.add_run("Performance Tuning: ").bold = True
skills.add_run("AWR, ASH, ADDM, SQL Tuning, Query Optimization\n")
skills.add_run("Automation & Scripting: ").bold = True
skills.add_run("Shell Scripting, Python, Crontab Scheduling\n")
skills.add_run("Cloud & OS: ").bold = True
skills.add_run("AWS (EC2, S3, RDS), Oracle Cloud Infrastructure (OCI), Linux (RHEL, CentOS)\n")
skills.add_run("High Availability & DR: ").bold = True
skills.add_run("RAC, Failover Configuration, Disaster Recovery (DR) Testing, GoldenGate Replication\n")

# Achievements
doc.add_heading("Achievements", level=2)
achievements = doc.add_paragraph(
    "• Ensured 99.9% database uptime through proactive performance tuning and monitoring.\n"
    "• Automated 6+ DBA tasks using Shell scripting, reducing manual efforts by 40%.\n"
    "• Improved SQL query response time by 35% via index optimization and AWR-based tuning.\n"
    "• Successfully performed multiple Oracle 19c RAC patch upgrades with zero production downtime.\n"
)

# Certifications
doc.add_heading("Certifications", level=2)
doc.add_paragraph("• CCNA (Cisco Certified Network Associate) – Besant Technologies")

# Education
doc.add_heading("Education", level=2)
doc.add_paragraph(
    "Bachelor of Technology (B.Tech.) – Information Technology\n"
    "PSN Engineering College | 2018 – 2022 | CGPA: 7.5"
)

# Save final optimized resume
optimized_path = "/workspaces/Devops/Peniel_Charles_Optimized_Oracle_DBA_Resume.docx"
doc.save(optimized_path)

optimized_path
